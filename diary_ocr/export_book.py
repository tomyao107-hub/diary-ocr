"""Chapter-aware book export (v1.3): Markdown / DOCX / PDF."""

from __future__ import annotations

import json
import re
from dataclasses import asdict, dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Iterable


EXPORT_SCHEMA_VERSION = 1


@dataclass
class Chapter:
    id: str
    title: str
    page_paths: list[str] = field(default_factory=list)  # relative paths
    notes: str = ""

    def to_dict(self) -> dict:
        return asdict(self)

    @classmethod
    def from_dict(cls, data: dict) -> "Chapter":
        return cls(
            id=str(data.get("id") or ""),
            title=str(data.get("title") or "未命名章节"),
            page_paths=[str(p) for p in data.get("page_paths", []) if isinstance(p, str)],
            notes=str(data.get("notes") or ""),
        )


@dataclass
class ExportProfile:
    title: str = "日记成册"
    author: str = ""
    include_toc: bool = True
    include_page_markers: bool = True
    cover_note: str = ""

    def to_dict(self) -> dict:
        return asdict(self)

    @classmethod
    def from_dict(cls, data: dict | None) -> "ExportProfile":
        data = data or {}
        return cls(
            title=str(data.get("title") or "日记成册"),
            author=str(data.get("author") or ""),
            include_toc=bool(data.get("include_toc", True)),
            include_page_markers=bool(data.get("include_page_markers", True)),
            cover_note=str(data.get("cover_note") or ""),
        )


@dataclass
class PageExportItem:
    path: str
    title: str
    text: str
    chapter_title: str | None = None
    missing: bool = False


@dataclass
class ExportCheck:
    ready: list[PageExportItem]
    missing: list[str]
    failed: list[str]


def load_page_text(image_path: str, output_dir: Path, image_paths: list[str], resolve_output) -> str | None:
    md_path = resolve_output(image_path, str(output_dir), image_paths)
    path = Path(md_path)
    if not path.exists():
        legacy = output_dir / f"{Path(image_path).stem}.md"
        path = legacy if legacy.exists() else path
    if not path.exists():
        return None
    try:
        return path.read_text(encoding="utf-8")
    except OSError:
        return None


def build_export_items(
    image_paths: list[str],
    output_dir: Path,
    resolve_output,
    chapters: list[Chapter] | None = None,
    project_root: Path | None = None,
) -> ExportCheck:
    ready: list[PageExportItem] = []
    missing: list[str] = []
    failed: list[str] = []

    path_to_chapter: dict[str, str] = {}
    if chapters and project_root is not None:
        for chapter in chapters:
            for rel in chapter.page_paths:
                try:
                    abs_path = str((project_root / rel).resolve())
                except (OSError, RuntimeError):
                    continue
                path_to_chapter[abs_path] = chapter.title

    for path in image_paths:
        text = load_page_text(path, output_dir, image_paths, resolve_output)
        name = Path(path).stem
        if text is None:
            missing.append(Path(path).name)
            continue
        if not text.strip():
            failed.append(Path(path).name)
        ready.append(
            PageExportItem(
                path=path,
                title=name,
                text=text,
                chapter_title=path_to_chapter.get(str(Path(path).resolve())),
            )
        )
    return ExportCheck(ready=ready, missing=missing, failed=failed)


def _group_by_chapter(items: list[PageExportItem]) -> list[tuple[str | None, list[PageExportItem]]]:
    if not any(item.chapter_title for item in items):
        return [(None, items)]
    groups: list[tuple[str | None, list[PageExportItem]]] = []
    current_title: str | None = object()  # type: ignore
    bucket: list[PageExportItem] = []
    for item in items:
        title = item.chapter_title
        if title != current_title:
            if bucket:
                groups.append((current_title if current_title is not object() else None, bucket))
            current_title = title
            bucket = [item]
        else:
            bucket.append(item)
    if bucket:
        groups.append((current_title if current_title is not object() else None, bucket))
    return groups


def export_markdown(
    items: list[PageExportItem],
    destination: Path,
    profile: ExportProfile,
) -> Path:
    destination.parent.mkdir(parents=True, exist_ok=True)
    lines: list[str] = [f"# {profile.title}", ""]
    if profile.author:
        lines.append(f"**作者：** {profile.author}")
        lines.append("")
    if profile.cover_note:
        lines.append(profile.cover_note)
        lines.append("")
    if profile.include_toc:
        lines.append("## 目录")
        lines.append("")
        for index, item in enumerate(items, 1):
            label = item.chapter_title or item.title
            lines.append(f"{index}. {label}")
        lines.append("")
    for chapter_title, group in _group_by_chapter(items):
        if chapter_title:
            lines.append(f"## {chapter_title}")
            lines.append("")
        for item in group:
            if profile.include_page_markers:
                lines.append(f"### {item.title}")
                lines.append("")
            lines.append(item.text.rstrip())
            lines.append("")
            lines.append("---")
            lines.append("")
    temporary = destination.with_suffix(destination.suffix + ".tmp")
    temporary.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")
    temporary.replace(destination)
    return destination


def export_docx(
    items: list[PageExportItem],
    destination: Path,
    profile: ExportProfile,
) -> Path:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as exc:
        raise RuntimeError(
            "导出 DOCX 需要 python-docx：pip install python-docx"
        ) from exc
    destination.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_heading(profile.title, level=0)
    if profile.author:
        document.add_paragraph(f"作者：{profile.author}")
    if profile.cover_note:
        document.add_paragraph(profile.cover_note)
    if profile.include_toc:
        document.add_heading("目录", level=1)
        for index, item in enumerate(items, 1):
            document.add_paragraph(
                f"{index}. {item.chapter_title or item.title}",
                style="List Number",
            )
    for chapter_title, group in _group_by_chapter(items):
        if chapter_title:
            document.add_heading(chapter_title, level=1)
        for item in group:
            if profile.include_page_markers:
                document.add_heading(item.title, level=2)
            for paragraph in item.text.splitlines() or [""]:
                run = document.add_paragraph(paragraph).runs
                if run:
                    run[0].font.size = Pt(12)
    document.save(str(destination))
    return destination


def export_pdf(
    items: list[PageExportItem],
    destination: Path,
    profile: ExportProfile,
) -> Path:
    """Export PDF via PyMuPDF (already a project dependency)."""
    try:
        import fitz
    except ImportError as exc:
        raise RuntimeError("导出 PDF 需要 PyMuPDF") from exc

    destination.parent.mkdir(parents=True, exist_ok=True)
    doc = fitz.open()
    width, height = fitz.paper_size("a4")
    margin = 54
    fontsize = 11
    line_height = fontsize * 1.5
    max_width = width - 2 * margin

    def new_page():
        page = doc.new_page(width=width, height=height)
        return page, margin

    def write_wrapped(page, y: float, text: str, size: float = fontsize, bold: bool = False):
        fontname = "china-s"  # built-in CJK fallback in PyMuPDF when available
        # Prefer a unicode font; fall back to helv for pure ASCII.
        try:
            # Insert text with textbox for wrapping.
            rect = fitz.Rect(margin, y, width - margin, height - margin)
            used = page.insert_textbox(
                rect,
                text,
                fontsize=size,
                fontname="helv",
                align=0,
            )
            # insert_textbox returns remaining unused height (negative if overflow).
            # Approximate consumed height.
            lines = max(1, text.count("\n") + 1)
            # Better: use text length estimate
            approx_lines = max(1, int(len(text) / 40) + text.count("\n") + 1)
            return y + approx_lines * size * 1.4
        except Exception:
            page.insert_text((margin, y + size), text[:200], fontsize=size, fontname="helv")
            return y + size * 1.5

    # Build plain text blocks for more reliable CJK rendering via TextWriter / HTML.
    # PyMuPDF Story is ideal when available.
    try:
        html_parts = [
            f"<h1>{_html_escape(profile.title)}</h1>",
        ]
        if profile.author:
            html_parts.append(f"<p><b>作者：</b>{_html_escape(profile.author)}</p>")
        if profile.cover_note:
            html_parts.append(f"<p>{_html_escape(profile.cover_note)}</p>")
        if profile.include_toc:
            html_parts.append("<h2>目录</h2><ol>")
            for item in items:
                html_parts.append(
                    f"<li>{_html_escape(item.chapter_title or item.title)}</li>"
                )
            html_parts.append("</ol>")
        for chapter_title, group in _group_by_chapter(items):
            if chapter_title:
                html_parts.append(f"<h2>{_html_escape(chapter_title)}</h2>")
            for item in group:
                if profile.include_page_markers:
                    html_parts.append(f"<h3>{_html_escape(item.title)}</h3>")
                for para in item.text.splitlines() or [""]:
                    html_parts.append(f"<p>{_html_escape(para)}</p>")
                html_parts.append("<hr/>")
        html = "\n".join(html_parts)
        story = fitz.Story(html)
        writer = fitz.DocumentWriter(str(destination))
        mediabox = fitz.paper_rect("a4")
        where = mediabox + (margin, margin, -margin, -margin)
        more = True
        while more:
            device = writer.begin_page(mediabox)
            more, _ = story.place(where)
            story.draw(device)
            writer.end_page()
        writer.close()
        doc.close()
        return destination
    except Exception:
        # Fallback: simple text pages (may not render CJK well without fonts).
        page, y = new_page()
        y = write_wrapped(page, y, profile.title, size=16)
        for item in items:
            block = f"\n[{item.title}]\n{item.text}\n"
            # naive pagination
            for line in block.splitlines():
                if y > height - margin - 20:
                    page, y = new_page()
                page.insert_text(
                    (margin, y),
                    line[:120],
                    fontsize=10,
                    fontname="helv",
                )
                y += 14
        doc.save(str(destination))
        doc.close()
        return destination


def _html_escape(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def save_export_profile(project_root: Path, profile: ExportProfile) -> Path:
    path = project_root / "output" / "exports" / "export_profile.json"
    path.parent.mkdir(parents=True, exist_ok=True)
    payload = {
        "schema_version": EXPORT_SCHEMA_VERSION,
        "saved_at": datetime.now().astimezone().isoformat(timespec="seconds"),
        "profile": profile.to_dict(),
    }
    temporary = path.with_suffix(path.suffix + ".tmp")
    temporary.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    temporary.replace(path)
    return path


def load_export_profile(project_root: Path) -> ExportProfile:
    path = project_root / "output" / "exports" / "export_profile.json"
    if not path.exists():
        return ExportProfile()
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
        return ExportProfile.from_dict(data.get("profile") if isinstance(data, dict) else {})
    except (OSError, json.JSONDecodeError):
        return ExportProfile()


def default_chapters_from_pages(relative_paths: Iterable[str]) -> list[Chapter]:
    return [
        Chapter(
            id="chapter-1",
            title="正文",
            page_paths=list(relative_paths),
        )
    ]
